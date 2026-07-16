"""Leitura de arquivos de entrada (atas, notas, documentos)."""

from __future__ import annotations

import csv
import io
from dataclasses import dataclass
from typing import BinaryIO

from docx import Document


EXTENSOES_SUPORTADAS = {".txt", ".md", ".markdown", ".docx", ".csv"}
TIPOS_UPLOAD = ["txt", "md", "docx", "csv"]


@dataclass
class DocumentoEntrada:
    nome: str
    texto: str


def _ler_texto_plano(raw: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _ler_docx(raw: bytes) -> str:
    doc = Document(io.BytesIO(raw))
    partes: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            partes.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                partes.append(" | ".join(cells))
    return "\n".join(partes).strip()


def _ler_csv(raw: bytes) -> str:
    texto = _ler_texto_plano(raw)
    leitor = csv.reader(io.StringIO(texto))
    linhas: list[str] = []
    for row in leitor:
        cells = [c.strip() for c in row if str(c).strip()]
        if cells:
            linhas.append(" | ".join(cells))
    return "\n".join(linhas).strip()


def extrair_texto_arquivo(nome: str, dados: bytes | BinaryIO) -> DocumentoEntrada:
    """Extrai texto de .txt, .md, .docx ou .csv."""
    if hasattr(dados, "read"):
        raw = dados.read()
    else:
        raw = dados

    nome_limpo = (nome or "documento").strip()
    lower = nome_limpo.lower()
    if lower.endswith(".docx"):
        texto = _ler_docx(raw)
    elif lower.endswith(".csv"):
        texto = _ler_csv(raw)
    elif lower.endswith((".txt", ".md", ".markdown")):
        texto = _ler_texto_plano(raw)
    else:
        raise ValueError(
            f"Formato não suportado: {nome_limpo}. Use .txt, .csv ou .docx."
        )

    if not texto.strip():
        raise ValueError(f"O arquivo '{nome_limpo}' está vazio ou sem texto legível.")

    return DocumentoEntrada(nome=nome_limpo, texto=texto.strip())


def montar_bloco_documentos(documentos: list[DocumentoEntrada]) -> str:
    """Concatena documentos em um bloco único para o prompt."""
    if not documentos:
        return ""
    blocos = []
    for i, doc in enumerate(documentos, start=1):
        blocos.append(f"--- Documento {i}: {doc.nome} ---\n{doc.texto}")
    return "\n\n".join(blocos)


def anexar_documento_sessao(
    documentos: list[dict],
    *,
    nome: str,
    texto: str,
) -> list[dict]:
    """Inclui ou substitui documento na lista de sessão (por nome)."""
    saida = [d for d in documentos if d.get("nome") != nome]
    saida.append({"nome": nome, "texto": texto})
    return saida
