"""Conversão pragmática de Markdown para Word (.docx)."""

from __future__ import annotations

import io
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor


_RE_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_RE_UL = re.compile(r"^[-*]\s+(.*)$")
_RE_OL = re.compile(r"^(\d+)[.)]\s+(.*)$")
_RE_BOLD = re.compile(r"\*\*(.+?)\*\*")
_RE_TABLE_SEP = re.compile(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")


def _eh_linha_tabela(linha: str) -> bool:
    s = linha.strip()
    return s.startswith("|") and s.count("|") >= 2


def _celulas(linha: str) -> list[str]:
    s = linha.strip().strip("|")
    return [c.strip() for c in s.split("|")]


def _adicionar_texto_com_negrito(paragrafo, texto: str) -> None:
    """Suporta **negrito** simples no meio do texto."""
    pos = 0
    for m in _RE_BOLD.finditer(texto):
        if m.start() > pos:
            paragrafo.add_run(texto[pos : m.start()])
        run = paragrafo.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(texto):
        paragrafo.add_run(texto[pos:])


def markdown_para_docx(texto: str, doc: Document) -> None:
    """Converte Markdown básico (headings, listas, tabelas, parágrafos) para o Document."""
    linhas = (texto or "").replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(linhas):
        linha = linhas[i]
        strip = linha.strip()

        if not strip:
            i += 1
            continue

        # Tabela Markdown
        if _eh_linha_tabela(strip):
            bloco: list[str] = []
            while i < len(linhas) and _eh_linha_tabela(linhas[i].strip()):
                if not _RE_TABLE_SEP.match(linhas[i].strip()):
                    bloco.append(linhas[i].strip())
                i += 1
            if bloco:
                cols = max(len(_celulas(r)) for r in bloco)
                table = doc.add_table(rows=len(bloco), cols=cols)
                table.style = "Table Grid"
                for r_idx, row_txt in enumerate(bloco):
                    cells = _celulas(row_txt)
                    for c_idx in range(cols):
                        cell_text = cells[c_idx] if c_idx < len(cells) else ""
                        cell = table.rows[r_idx].cells[c_idx]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        _adicionar_texto_com_negrito(p, cell_text)
                        if r_idx == 0:
                            for run in p.runs:
                                run.bold = True
            continue

        m_h = _RE_HEADING.match(strip)
        if m_h:
            nivel = min(len(m_h.group(1)), 3)
            doc.add_heading(m_h.group(2).strip(), level=nivel)
            i += 1
            continue

        m_ul = _RE_UL.match(strip)
        if m_ul:
            p = doc.add_paragraph(style="List Bullet")
            _adicionar_texto_com_negrito(p, m_ul.group(1))
            i += 1
            continue

        m_ol = _RE_OL.match(strip)
        if m_ol:
            p = doc.add_paragraph(style="List Number")
            _adicionar_texto_com_negrito(p, m_ol.group(2))
            i += 1
            continue

        p = doc.add_paragraph()
        _adicionar_texto_com_negrito(p, strip)
        i += 1


def criar_documento(titulo: str) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    heading = doc.add_heading(titulo, level=0)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x00, 0x10, 0x60)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return doc


def markdown_para_docx_bytes(titulo: str, markdown: str) -> bytes:
    """Converte Markdown em .docx e retorna os bytes."""
    doc = criar_documento(titulo)
    markdown_para_docx(markdown, doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def salvar_markdown_como_docx(caminho: Path, titulo: str, markdown: str) -> Path:
    """Cria um .docx a partir de um texto Markdown e salva em `caminho`."""
    caminho.parent.mkdir(parents=True, exist_ok=True)
    caminho.write_bytes(markdown_para_docx_bytes(titulo, markdown))
    return caminho
